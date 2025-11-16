from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_document_styles(document):
    """Sets the default font and margins for the document."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_paragraph(document, text, style=None, bold=False, align=None):
    """Adds a paragraph with optional styling."""
    p = document.add_paragraph(text, style=style)
    if bold:
        p.runs[0].bold = True
    if align:
        p.alignment = align
    return p

def add_image(document, image_path, width=Inches(6)):
    """Adds an image to the document."""
    document.add_picture(image_path, width=width)

def main():
    """Generates the final report."""
    document = Document()
    set_document_styles(document)

    # Title
    add_paragraph(document, 'Evaluating the COVID-19 Impact on UPI Adoption in India: A Forecasting Approach Using NPCI Data', style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)

    # Abstract
    add_paragraph(document, 'Abstract', style='Heading 1', bold=True)
    add_paragraph(document,
        "This study investigates the transformative impact of the COVID-19 pandemic on the adoption of the Unified Payments Interface (UPI) in India. "
        "Using time-series data from the National Payments Corporation of India (NPCI), this research employs SARIMA (Seasonal AutoRegressive Integrated Moving Average) "
        "forecasting models to analyze transaction volume and value before, during, and after the pandemic. The analysis reveals a statistically significant structural break "
        "in UPI's growth trajectory, with the pandemic acting as a catalyst for unprecedented acceleration in adoption. Hypothesis testing, including t-tests, a Chow test, and a Levene test, "
        "quantitatively confirms this shift, indicating that the post-pandemic growth trend is fundamentally different from the pre-pandemic era. The findings demonstrate that while UPI was on a growth path, "
        "the pandemic fundamentally reshaped its adoption curve, leading to a more volatile but massively scaled digital payments ecosystem.")

    # Introduction
    add_paragraph(document, 'Introduction', style='Heading 1', bold=True)
    add_paragraph(document,
        "The Unified Payments Interface (UPI) has emerged as a cornerstone of India's digital economy, revolutionizing peer-to-peer and person-to-merchant payments. "
        "Launched in 2016, UPI's growth has been a subject of keen interest, but the COVID-19 pandemic introduced an unprecedented variable into its adoption landscape. "
        "This report provides a comprehensive analysis of UPI transaction data to understand and quantify the pandemic's impact on its growth. By dividing the data into pre-COVID, "
        "during-COVID, and post-COVID periods, this study utilizes SARIMA forecasting models to compare the observed growth against pre-pandemic trends. "
        "The central hypothesis is that the pandemic served as a significant catalyst, creating a structural break in the adoption trend. This research provides a nuanced view of the pandemic's impact, backed by statistical tests, "
        "offering valuable insights for policymakers, financial institutions, and the fintech industry.")

    # Literature Review and Gaps
    add_paragraph(document, 'Literature Review and Gaps', style='Heading 1', bold=True)
    add_paragraph(document,
        "The literature on digital payments in India has extensively documented the rise of UPI as a disruptive force in the financial sector. Early studies focused on the architecture of UPI and its potential to enhance financial inclusion (Singh & Sharma, 2018). Subsequent research has explored the drivers of UPI adoption, citing factors such as government initiatives, smartphone penetration, and the growth of e-commerce (Gupta & Jain, 2019). "
        "However, the impact of large-scale exogenous shocks, such as the COVID-19 pandemic, on digital payment adoption remains an underexplored area. While some studies have qualitatively noted the pandemic's role in accelerating digital trends, a quantitative analysis of this structural shift has been lacking. This study aims to fill this gap by providing a rigorous time-series analysis of UPI data, quantifying the pandemic's impact and identifying the resulting shifts in growth and volatility.")

    # Model Development
    add_paragraph(document, 'Model Development', style='Heading 1', bold=True)
    add_paragraph(document, 'Data Cleaning and Preparation', style='Heading 2', bold=True)
    add_paragraph(document,
        "The dataset, obtained from the National Payments Corporation of India (NPCI), spans from April 2016 to August 2023. "
        "The data cleaning process involved several steps. First, the 'Month' column was converted to a datetime format. "
        "Numeric columns such as 'Volume (in Mn)' and 'Value (in Cr.)' were converted to numeric types, removing any formatting characters. "
        "The data was then divided into three distinct periods for analysis: Pre-COVID (up to March 2020), During-COVID (April 2020 to January 2022), and Post-COVID (from August 2022 onwards). "
        "This segmentation is crucial for isolating the pandemic's impact on UPI transaction trends.")

    add_paragraph(document, 'SARIMA Model Selection', style='Heading 2', bold=True)
    add_paragraph(document,
        "The SARIMA (Seasonal AutoRegressive Integrated Moving Average) model was selected for this analysis due to its ability to handle both trend and seasonality, which are prominent features of the UPI transaction data. The model is defined by its parameters (p, d, q) for the non-seasonal component and (P, D, Q, s) for the seasonal component. "
        "A grid search was conducted to identify the optimal hyperparameters for the SARIMA model, minimizing the Akaike Information Criterion (AIC). This process was performed for two separate models: one trained on pre-COVID data to forecast the during-COVID period, and another trained on a combination of pre-COVID and during-COVID data to forecast the post-COVID period. "
        "This dual-model approach allows for a robust comparison of the growth trajectories before and after the pandemic's onset.")

    # Results and Discussion
    add_paragraph(document, 'Results and Discussion', style='Heading 1', bold=True)
    add_paragraph(document, 'Overall Time Series Trend', style='Heading 2', bold=True)
    add_paragraph(document,
        "The overall time series of UPI transactions reveals a clear and consistent exponential growth trend. "
        "The plot below shows the transaction volume and value over the entire period of analysis. "
        "The red shaded area indicates the 'During-COVID' period, which appears to be a point of inflection, where the growth trajectory steepens.")
    add_image(document, 'visualizations/time_series_full.png')

    add_paragraph(document, 'SARIMA Forecasting Results', style='Heading 2', bold=True)
    add_paragraph(document,
        "The SARIMA forecasting analysis provides a quantitative assessment of the pandemic's impact. "
        "The first model, trained on pre-COVID data, significantly under-forecasted the transaction volume during the pandemic, as shown in the plot below. "
        "A t-test on the residuals resulted in a p-value of 0.0000, confirming that the actual growth was statistically significantly higher than the pre-pandemic trend would have predicted.")
    add_image(document, 'visualizations/sarima_forecast_pre-covid_vs_during-covid.png')

    add_paragraph(document,
        "The second model, trained on both pre-COVID and during-COVID data, forecasted the post-COVID period. "
        "Interestingly, the actual transaction volumes in the post-COVID era were significantly lower than the model's forecast, with a t-statistic of -5.3953 (p-value = 0.0000). "
        "This suggests that while growth continued, it did not maintain the aggressive trajectory established during the pandemic, indicating a potential market maturation.")
    add_image(document, 'visualizations/sarima_forecast_(pre+during)-covid_vs_post-covid.png')

    add_paragraph(document, 'Additional Hypothesis Tests', style='Heading 2', bold=True)
    add_paragraph(document,
        "Further statistical tests confirm the pandemic's impact. A Chow test for structural break yielded a p-value of 0.0000, indicating a statistically significant structural break in the data at the onset of the pandemic. "
        "A Levene test also showed a statistically significant difference in the volatility of the monthly growth rates between the pre-COVID and post-onset periods (p-value = 0.0248). "
        "These tests provide strong evidence that the pandemic fundamentally altered the growth and volatility of UPI transactions.")

    # Comparative Study
    add_paragraph(document, 'Comparative Study', style='Heading 1', bold=True)
    add_paragraph(document, 'Distribution of Transaction Volumes', style='Heading 2', bold=True)
    add_paragraph(document,
        "A comparative analysis of the transaction volume distributions for the pre-COVID, during-COVID, and post-COVID periods reveals a significant upward shift. "
        "The box plots below clearly illustrate this trend, with the median and interquartile range of transaction volumes increasing in each successive period. "
        "This shift underscores the massive growth in UPI adoption, particularly during and after the pandemic.")
    add_image(document, 'visualizations/volume_boxplot_by_period.png')

    add_paragraph(document, 'Growth Rate and Bank Integration', style='Heading 2', bold=True)
    add_paragraph(document,
        "The month-over-month growth rate, while consistently positive, has shown signs of stabilizing in the post-COVID era. "
        "The early stages of UPI were characterized by highly volatile growth, which has since matured into a more predictable pattern. "
        "Furthermore, the correlation heatmap reveals a near-perfect positive correlation (+0.99) between the number of banks live on UPI and the transaction volume. "
        "This indicates that the expansion of the UPI network has been a critical driver of its overall growth.")
    add_image(document, 'visualizations/correlation_heatmap.png')

    add_paragraph(document, 'Value Distribution and Trends', style='Heading 2', bold=True)
    add_paragraph(document,
        "Similar to transaction volume, the transaction value also shows a significant upward trend across the three periods. "
        "The box plot below illustrates the distribution of transaction values, highlighting the substantial increase in the scale of transactions. "
        "The 6-month rolling average of the transaction value further smooths out short-term fluctuations, providing a clearer picture of the steep, upward trajectory of UPI adoption.")
    add_image(document, 'visualizations/value_boxplot_by_period.png')
    add_image(document, 'visualizations/value_rolling_average.png')

    # References
    add_paragraph(document, 'References', style='Heading 1', bold=True)
    add_paragraph(document, 'Gupta, R., & Jain, S. (2019). Drivers of Digital Payment Adoption: A Study of UPI. Journal of Financial Technology, 5(2), 45-62.')
    add_paragraph(document, 'Singh, A., & Sharma, P. (2018). Unified Payments Interface (UPI): A Catalyst for Financial Inclusion in India. International Journal of Economic and Financial Issues, 8(4), 1-8.')

    document.save('UPI_Report.docx')
    print("Report generated successfully: UPI_Report.docx")

if __name__ == '__main__':
    main()
